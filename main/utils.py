# ===================== STANDARD IMPORTS =====================
import os
import json

# ===================== FILE PARSING =====================
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

# ===================== NLP & ML =====================
import spacy
from transformers import pipeline

# ===================== GEMINI (NEW SDK) =====================
from google import genai
from django.conf import settings

# ============================================================
# INITIALIZATIONS
# ============================================================

# Summarizer
summarizer = pipeline(
    "summarization",
    model="sshleifer/distilbart-cnn-12-6"
)

# spaCy NLP
nlp = spacy.load("en_core_web_sm")

# Gemini client (Fixed model string and added safety)
client = genai.Client(api_key=settings.GOOGLE_API_KEY)

# Risk keywords
PENALTY_TERMS = [
    "liable", "penalty", "breach", "termination", "fine",
    "indemnify", "compensation", "non-compliance", "forfeit", "void"
]


# ============================================================
# TEXT EXTRACTION
# ============================================================

def extract_text_from_file(file):
    """Extract text from Django uploaded file."""
    text = ""
    filename = file.name.lower()

    if filename.endswith(".pdf"):
        # Reset file pointer to beginning
        file.seek(0)
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()

    elif filename.endswith(".docx"):
        doc = Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"

    elif filename.endswith((".jpg", ".jpeg", ".png")):
        image = Image.open(file)
        text = pytesseract.image_to_string(image)

    else:
        file.seek(0)
        text = file.read().decode("utf-8")

    return text.strip()


def extract_text_from_path(file_path):
    """Extract text from stored file path."""
    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".jpg", ".jpeg", ".png"]:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)

    elif ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

    return text.strip()


# ============================================================
# GEMINI DOCUMENT UNDERSTANDING
# ============================================================

def extract_legal_entities_with_gemini(document_text):
    """
    Evidence-based extraction.
    Gemini only reports WHAT IT FINDS, not whether it is compliant.
    """

    prompt = f"""
You are an Indian legal document evidence extractor.

For each requirement listed below:
- List ALL evidence found in the document
- Mention wording used
- Mention approximate section/page if possible
- If unclear, say "uncertain"
- If not found, return empty list

Return STRICT JSON only.

FORMAT:
{{
  "execution_date": [{{"value": "", "evidence": "", "confidence": ""}}],
  "execution_place": [],
  "licensor_names": [],
  "licensee_names": [],
  "stamp_registration": [],
  "witnesses": [],
  "signatures": []
}}

IMPORTANT:
- Understand Indian legal terms like:
  "HEREINAFTER called", "Registered as Document No.",
  "Witness of execution", "Thumb impression verified",
  "Digitally signed", "Schedule", "GRN", "UIDAI"

DOCUMENT:
\"\"\"{document_text[:12000]}\"\"\"
"""

    try:
        response = client.models.generate_content(
            model="models/gemini-1.0-pro",
            contents=prompt
        )

        content = response.text
        if "```" in content:
            content = content.split("```")[1]

        return json.loads(content.strip())

    except Exception as e:
        print("Gemini extraction error:", e)
        return {}


# ============================================================
# SUMMARIZATION
# ============================================================

def summarize_text(text):
    if not text: return "No text found."
    if len(text) < 1000:
        return summarizer(text, max_length=100, min_length=30, do_sample=False)[0]["summary_text"]

    chunks = [text[i:i + 1000] for i in range(0, len(text), 1000)]
    summary = ""
    for chunk in chunks[:3]: # Limit chunks for performance
        part = summarizer(chunk, max_length=100, min_length=30, do_sample=False)[0]["summary_text"]
        summary += " " + part
    return summary.strip()


def find_risky_clauses(text):
    doc = nlp(text[:10000])
    return [sent.text for sent in doc.sents if any(term in sent.text.lower() for term in PENALTY_TERMS)]

# doc comparisons
import re
from PIL import Image
import pytesseract
from io import BytesIO

def extract_text_from_file(uploaded_file):
    """
    OCR without saving file to disk
    """
    image = Image.open(BytesIO(uploaded_file.read()))
    text = pytesseract.image_to_string(image)
    return normalize_text(text)

def normalize_text(text):
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = text.replace(' ,', ',').replace(' .', '.')
    return text.strip()

import re
from io import BytesIO
from PIL import Image
import pytesseract
import pdfplumber
import docx

def extract_text_from_file(uploaded_file):
    file_name = uploaded_file.name.lower()

    # ---------- IMAGE ----------
    if file_name.endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(BytesIO(uploaded_file.read()))
        text = pytesseract.image_to_string(image)

    # ---------- PDF ----------
    elif file_name.endswith('.pdf'):
        text = extract_text_from_pdf(uploaded_file)

    # ---------- DOCX ----------
    elif file_name.endswith('.docx'):
        doc = docx.Document(uploaded_file)
        text = "\n".join(p.text for p in doc.paragraphs)

    else:
        raise ValueError("Unsupported file format")

    return normalize_text(text)


def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    # If text-based PDF failed → fallback to OCR
    if not text.strip():
        uploaded_file.seek(0)
        images = convert_pdf_to_images(uploaded_file) # type: ignore
        for img in images:
            text += pytesseract.image_to_string(img)

    return text


def normalize_text(text):
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    return text.strip()
